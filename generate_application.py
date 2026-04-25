#!/usr/bin/env python3
"""
generate_application.py — Offline CV + cover letter generator.

Reads  : config/job_description.md   (job metadata + description)
         config/profile.json          (candidate master CV)

Writes : cv/<company>_<title>/
             <name>_CV.pdf            (tailored CV)
             <name>_Cover_Letter.docx (cover letter)
             <name>_Career_Filler.txt (form-filler data)
         cv/tracker.xlsx              (application log — status updated manually)

job_description.md format
--------------------------
    ---
    title: Senior Data Analyst
    company: Acme Corp
    location: Remote
    url: https://linkedin.com/jobs/view/12345/
    ---

    Full job description text here...

Usage:
    python generate_application.py
"""

import json
import os
import re
from datetime import datetime
from pathlib import Path

import anthropic
from docx import Document
from docx.shared import Cm, Pt, RGBColor
from dotenv import load_dotenv

from cv_generator import generate_tailored_cv, render_cv_pdf
import career_education_filler as cef

load_dotenv()

CONFIG_PATH = Path("config/profile.json")
JD_PATH = Path("config/job_description.md")
CV_DIR = Path("cv")
TRACKER_PATH = CV_DIR / "tracker.xlsx"

BRAND_BLUE = RGBColor(0x0A, 0x66, 0xC2)
DARK = RGBColor(0x1B, 0x1B, 0x1B)
MID_GREY = RGBColor(0x55, 0x55, 0x55)


# ---------------------------------------------------------------------------
# Loaders
# ---------------------------------------------------------------------------


def load_profile() -> dict:
    with open(CONFIG_PATH, encoding="utf-8") as f:
        return json.load(f)


def parse_job_description() -> tuple[dict, str]:
    """Parse metadata and description text from config/job_description.md."""
    text = JD_PATH.read_text(encoding="utf-8")

    job: dict = {
        "title": "",
        "company": "",
        "location": "",
        "url": "",
        "easy_apply": False,
        "scraped_at": datetime.now().strftime("%Y-%m-%d %H:%M"),
    }

    match = re.match(r"^---\s*\n(.*?)\n---\s*\n(.*)", text, re.DOTALL)
    if match:
        meta_block, description = match.group(1), match.group(2).strip()
        for line in meta_block.splitlines():
            if ":" in line:
                key, _, val = line.partition(":")
                key = key.strip().lower()
                val = val.strip()
                if key in job:
                    job[key] = val
    else:
        description = text.strip()

    if not job["title"]:
        job["title"] = JD_PATH.stem.replace("_", " ").title()
    if not job["company"]:
        job["company"] = "Unknown Company"

    return job, description


# ---------------------------------------------------------------------------
# Cover letter — Claude generation
# ---------------------------------------------------------------------------


def generate_cover_letter(
    profile: dict, job: dict, description: str, tailored_cv: dict
) -> str:
    """Call Claude Opus to write a cover letter; returns plain paragraphs."""
    api_key = os.environ.get("API_KEY") or os.environ.get("ANTHROPIC_API_KEY")
    client = anthropic.Anthropic(api_key=api_key)

    personal = profile["personal"]

    prompt = f"""Write a professional cover letter for this candidate applying to the role below.

CANDIDATE:
Name     : {personal['full_name']}
Email    : {personal['email']}
Location : {personal['location']}
Summary  : {tailored_cv.get('summary', '')}
Top skills: {', '.join(tailored_cv.get('skills', [])[:6])}

TARGET ROLE:
Title   : {job['title']}
Company : {job['company']}
Location: {job.get('location', '')}

JOB DESCRIPTION:
{description or '(No description available — tailor to the job title and company.)'}

INSTRUCTIONS:
- Write exactly 3 paragraphs separated by a blank line
- Opening: genuine, specific interest in THIS role and company
- Middle: connect 2–3 concrete achievements from the CV directly to key job requirements
- Closing: confident call-to-action inviting next steps
- Tone: professional but natural — confident and direct, not stiff
- Do NOT include date, address blocks, salutation, or sign-off — body paragraphs only
- Return plain text, no markdown, no headers, no bullet points"""

    resp = client.messages.create(
        model="claude-opus-4-6",
        max_tokens=800,
        messages=[{"role": "user", "content": prompt}],
    )
    return resp.content[0].text.strip()


# ---------------------------------------------------------------------------
# Cover letter — Word rendering
# ---------------------------------------------------------------------------


def _add_run(para, text: str, size: int, bold: bool = False, color: RGBColor = DARK):
    run = para.add_run(text)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = color
    return run


def save_cover_letter_docx(
    cover_letter_text: str,
    profile: dict,
    job: dict,
    output_path: Path,
) -> None:
    """Render the cover letter as a branded Word document."""
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.0)

    personal = profile["personal"]

    # --- Candidate name ---
    name_para = doc.add_paragraph()
    _add_run(name_para, personal["full_name"], size=18, bold=True, color=BRAND_BLUE)
    name_para.paragraph_format.space_after = Pt(2)

    # --- Contact line ---
    contact_parts = [
        personal.get("email", ""),
        personal.get("phone", ""),
        personal.get("location", ""),
    ]
    if personal.get("linkedin"):
        contact_parts.append(personal["linkedin"])
    contact_para = doc.add_paragraph()
    _add_run(
        contact_para,
        "  |  ".join(p for p in contact_parts if p),
        size=9,
        color=MID_GREY,
    )
    contact_para.paragraph_format.space_after = Pt(4)

    # --- Thin divider ---
    hr = doc.add_paragraph()
    _add_run(hr, "─" * 90, size=7, color=BRAND_BLUE)
    hr.paragraph_format.space_after = Pt(10)

    # --- Date ---
    date_para = doc.add_paragraph()
    _add_run(date_para, datetime.now().strftime("%B %d, %Y"), size=10, color=MID_GREY)
    date_para.paragraph_format.space_after = Pt(14)

    # --- Salutation ---
    sal = doc.add_paragraph()
    _add_run(sal, f"Dear Hiring Manager at {job['company']},", size=11, bold=True)
    sal.paragraph_format.space_after = Pt(10)

    # --- Body paragraphs ---
    for para_text in cover_letter_text.split("\n\n"):
        para_text = para_text.strip()
        if not para_text:
            continue
        p = doc.add_paragraph()
        _add_run(p, para_text, size=11)
        p.paragraph_format.space_after = Pt(10)
        p.paragraph_format.line_spacing = Pt(15)

    # --- Closing ---
    closing = doc.add_paragraph()
    _add_run(closing, "Sincerely,", size=11)
    closing.paragraph_format.space_before = Pt(10)
    closing.paragraph_format.space_after = Pt(22)

    sig = doc.add_paragraph()
    _add_run(sig, personal["full_name"], size=11, bold=True, color=BRAND_BLUE)

    doc.save(str(output_path))
    print(f"    [cover] Saved -> {output_path.resolve()}")


# ---------------------------------------------------------------------------
# Career form filler
# ---------------------------------------------------------------------------


def save_career_filler(profile: dict, output_path: Path) -> None:
    """Generate career filler content using career_education_filler and save to output_path."""
    personal = profile.get("personal", {})
    cv = profile.get("cv", {})

    header = "\n".join([
        cef._rule("═", 60),
        "APPLICATION FORM FILLER",
        f"Generated : {datetime.now().strftime('%Y-%m-%d %H:%M')}",
        f"Candidate : {personal.get('full_name', '')}",
        f"Email     : {personal.get('email', '')}",
        f"Phone     : {personal.get('phone', '')}",
        f"Location  : {personal.get('location', '')}",
        f"LinkedIn  : {personal.get('linkedin', '')}",
        cef._rule("═", 60),
        "",
    ])

    sections = [
        header,
        cef.format_experience(cv.get("experience", [])),
        cef.format_education(cv.get("education", [])),
        cef.format_skills(cv.get("skills", [])),
        cef.format_certifications(cv.get("certifications", [])),
        cef.format_languages(cv.get("languages", [])),
    ]

    content = "\n".join(s for s in sections if s)
    output_path.write_text(content, encoding="utf-8")
    print(f"    [filler] Saved -> {output_path.resolve()}")


# ---------------------------------------------------------------------------
# Application tracker
# ---------------------------------------------------------------------------


def update_tracker(job: dict, job_folder: Path) -> None:
    """Append a new row to cv/tracker.xlsx; creates the file with headers if absent."""
    try:
        from openpyxl import load_workbook
        wb = load_workbook(TRACKER_PATH)
        ws = wb.active
    except FileNotFoundError:
        from openpyxl import Workbook
        wb = Workbook()
        ws = wb.active
        ws.append(["Date", "Company", "Job Title", "URL", "Folder", "Status"])

    ws.append([
        datetime.now().strftime("%Y-%m-%d %H:%M"),
        job["company"],
        job["title"],
        job.get("url", ""),
        str(job_folder.resolve()),
        "",  # status — update manually
    ])

    wb.save(TRACKER_PATH)
    print(f"    [tracker] Entry added -> {TRACKER_PATH.resolve()}")


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------


def main() -> None:
    print("=" * 60)
    print("Application Generator")
    print("=" * 60)

    if not JD_PATH.exists():
        print(f"[!] {JD_PATH} not found.")
        print("    Create it with the format described in the script docstring.")
        return

    profile = load_profile()
    job, description = parse_job_description()

    print(f"Role    : {job['title']}")
    print(f"Company : {job['company']}")
    print(f"Desc    : {len(description)} chars")
    print()

    # Build paths
    company_slug = re.sub(r"[^\w]+", "_", job["company"]).strip("_")[:30]
    title_slug = re.sub(r"[^\w]+", "_", job["title"]).strip("_")[:30]
    name_slug = "_".join(profile["personal"].get("full_name", "Candidate").split())

    job_folder = CV_DIR / f"{company_slug}_{title_slug}"
    job_folder.mkdir(parents=True, exist_ok=True)

    cv_path = job_folder / f"{name_slug}_CV.pdf"
    cl_path = job_folder / f"{name_slug}_Cover_Letter.docx"
    filler_path = job_folder / f"{name_slug}_Career_Filler.txt"

    # --- Step 1: Tailored CV dict ---
    print("[1/4] Calling Claude to tailor CV...")
    tailored_cv = generate_tailored_cv(profile, job, description)

    # --- Step 2: CV PDF ---
    print("[2/4] Rendering CV PDF...")
    render_cv_pdf(tailored_cv, profile, job, cv_path)

    # --- Step 3: Cover letter ---
    print("[3/4] Generating cover letter...")
    cover_letter_text = generate_cover_letter(profile, job, description, tailored_cv)
    save_cover_letter_docx(cover_letter_text, profile, job, cl_path)

    # --- Step 4: Career form filler ---
    print("[4/4] Saving career form filler...")
    save_career_filler(profile, filler_path)

    # --- Tracker ---
    print("[ + ] Updating application tracker...")
    update_tracker(job, job_folder)

    print()
    print("Done.")
    print(f"  Folder      : {job_folder.resolve()}")
    print(f"  CV          : {cv_path.name}")
    print(f"  Cover letter: {cl_path.name}")
    print(f"  Form filler : {filler_path.name}")
    print(f"  Tracker     : {TRACKER_PATH.resolve()}")


if __name__ == "__main__":
    main()
